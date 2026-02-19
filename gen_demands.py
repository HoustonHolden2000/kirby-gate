"""Generate demand letters for all delinquent Kirby Gate targets from live database."""

import os
import sqlite3
import sys
from datetime import datetime, timedelta

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASEDIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASEDIR, "kirbygate.db")

TOTAL_CAMPUS_SQFT = 672_718
CURRENT_WEEKLY_RATE = 9_000.00
DECLARATION_DATE = "May 5, 2011"
LIEN_DEADLINE = "April 1, 2026"
CURE_DAYS = 30
TODAY_STR = datetime.now().strftime("%B %d, %Y")
CURE_DATE = (datetime.now() + timedelta(days=CURE_DAYS)).strftime("%B %d, %Y")

INSTRUMENT_CITATION = (
    "This notice is issued pursuant to Section 5 (Default) and Section 6 "
    "(Remedies) of the Declaration of Restrictive Covenants, Instrument "
    "No. 12106392, recorded May 5, 2011, Office of the Register of Shelby "
    "County, Tennessee."
)

# ── Per-entity contact salutations (correction #7) ─────────────────────────
# Key = business_name prefix match; value = salutation line
# Salutations updated per Shelby County Assessor verified owners (Tax Year 2025)
CONTACTS = {
    "Kroger": "Dear Kroger Limited Partnership I, Property Tax Department:",
    "Starbucks": "Dear Trustee, Kristovich Trust:",
    "Wendy": "Dear Arline Townhomes LLC:",
    "Summit of Germantown": "Dear Memphis Senior Housing Propco K6 LLC:",
    "Freedom Plasma": "Dear Realty Income Properties 25 LLC:",
    "Dollar General": "Dear DG Memphis LLC:",
    "Dunkin": "Dear Kirby Land Holdings LLC:",
    "Pointe at Kirby": "Dear 6480 Quince Road East Holdings LLC:",
    "Dollar Tree": "Dear DT Retail Properties LLC:",
    "KG Business": "Dear 3LS Properties Inc:",
}

# Mailing addresses verified via Shelby County GIS CERT_Parcel layer (2025)
CONTACT_ADDRESSES = {
    "Kroger": "Kroger Limited Partnership I\nProperty Tax, 7th Floor\n1014 Vine St\nCincinnati, OH 45202",
    "Starbucks": "Kristovich Trust\n21555 Prospect Rd\nSaratoga, CA 95070",
    "Wendy": "Arline Townhomes LLC\nc/o Deloris Mayuga\n6571 El Roble St\nLong Beach, CA 90815",
    "Summit of Germantown": "Memphis Senior Housing Propco K6 LLC\nc/o Kayne Anderson Real Estate Advisors\n1 Town Center Rd, Ste 300\nBoca Raton, FL 33486",
    "Freedom Plasma": "Realty Income Properties 25 LLC\n1430 E Southlake Blvd, Suite 200\nSouthlake, TX 76092",
    "Dollar General": "DG Memphis LLC\n12155 Spencer Rd\nMilford, MI 48380",
    "Dunkin": "Kirby Land Holdings LLC\n1050 Cambridge Sq\nAlpharetta, GA 30009",
    "Pointe at Kirby": "6480 Quince Road East Holdings LLC\nc/o CWCapital Asset Management LLC\n900 19th St NW, Floor 8\nWashington, DC 20006",
    "Dollar Tree": "DT Retail Properties LLC\n500 Volvo Parkway\nChesapeake, VA 23320",
    "KG Business": "3LS Properties Inc\n301 S Perimeter Dr, Suite 200\nNashville, TN 37211",
}


def get_contact_salutation(business_name):
    for key, salutation in CONTACTS.items():
        if business_name.startswith(key):
            return salutation
    return "Dear Sir or Madam:"


def get_contact_address_block(business_name):
    for key, addr in CONTACT_ADDRESSES.items():
        if business_name.startswith(key):
            return addr
    return None


def money(val):
    if val is None:
        return "TBD"
    return f"${val:,.2f}"


def build_letter(row):
    sqft = row["sqft"] or 0
    pct = sqft / TOTAL_CAMPUS_SQFT
    past_due = row["past_due_balance"]
    weekly = row["weekly_rate"] or (CURRENT_WEEKLY_RATE * pct)
    fwd_monthly = weekly * 4.333

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Letterhead ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VANGUARD SECURITY SERVICES")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Designated Agent for Kirby Gate Security Operations")

    # Refinement #1: Phone and email placeholders
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Phone: [INSERT]    Email: [INSERT]")

    doc.add_paragraph("")
    doc.add_paragraph(TODAY_STR)
    doc.add_paragraph("")

    # ── Addressee ──
    doc.add_paragraph("SENT VIA CERTIFIED MAIL")
    # Refinement #5: Certified mail tracking blank
    doc.add_paragraph("Certified Mail No.: _______________")
    doc.add_paragraph("")

    # Correction #7: specific contact address block
    contact_block = get_contact_address_block(row["business_name"])
    if contact_block:
        for line in contact_block.split("\n"):
            doc.add_paragraph(line)

    # Correction #4: Remove duplicate RE: — only show property/tenant info
    doc.add_paragraph(f"Property: {row['address']}")
    doc.add_paragraph(f"Tenant/Occupant: {row['business_name']}")
    if row["entity_owner"]:
        doc.add_paragraph(f"Entity: {row['entity_owner']}")
    doc.add_paragraph("")

    # Subject line (single RE:)
    p = doc.add_paragraph()
    run = p.add_run("RE: NOTICE OF COVENANT NON-COMPLIANCE AND DEMAND FOR CURE")
    run.bold = True
    run.underline = True
    doc.add_paragraph("")

    # ── Correction #7: Specific contact salutation ──
    salutation = get_contact_salutation(row["business_name"])
    doc.add_paragraph(salutation)
    doc.add_paragraph("")

    doc.add_paragraph(
        "This letter constitutes formal notice that the above-referenced property "
        "is in material default of the Declaration of Restrictive Covenants recorded "
        f'on {DECLARATION_DATE} (the "Declaration") governing the Kirby Gate '
        "commercial development, Memphis, Tennessee."
    )
    doc.add_paragraph("")

    # ── Updated designation language ──
    doc.add_paragraph(
        "Pursuant to the Declaration, each parcel owner is obligated to fund its "
        "pro-rata share of campus-wide security services. Vanguard Security Services "
        "has been designated as the authorized security provider under the Declaration, "
        "as confirmed by the Security Management Agreement dated February 12, 2026, "
        "appointing Vanguard Security as designated agent for enforcement of security "
        "charges under the Declaration."
    )
    doc.add_paragraph("")

    # Refinement #2: Compliance context paragraph
    doc.add_paragraph(
        "For your reference, approximately half of all parcel owners within "
        "Kirby Gate have maintained continuous compliance with these security "
        "obligations since the Declaration was recorded. The above-referenced "
        "property is among those parcels that have not fulfilled this obligation."
    )
    doc.add_paragraph("")

    # Arrears calculation
    p = doc.add_paragraph()
    run = p.add_run("ARREARS CALCULATION:")
    run.bold = True

    doc.add_paragraph(f"  Total Campus Square Footage:     {TOTAL_CAMPUS_SQFT:,} SF")
    doc.add_paragraph(f"  Your Parcel Square Footage:      {sqft:,} SF")
    doc.add_paragraph(f"  Your Pro-Rata Share:             {pct:.4%}")
    doc.add_paragraph(f"  Your Weekly Rate:                {money(weekly)}/week")
    doc.add_paragraph(f"  Arrears Period:                  36 months (156 weeks)")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    if past_due is not None and past_due > 0:
        run = p.add_run(f"  TOTAL 36-MONTH ARREARS OWED:     {money(past_due)}")
    else:
        run = p.add_run(
            "  TOTAL 36-MONTH ARREARS OWED:     "
            "TO BE DETERMINED (pending reconciliation)"
        )
    run.bold = True
    doc.add_paragraph("")

    # Forward billing
    p = doc.add_paragraph()
    run = p.add_run("FORWARD BILLING (Effective January 2026):")
    run.bold = True
    doc.add_paragraph(f"  Current Weekly Rate:             {money(weekly)}/week")
    doc.add_paragraph(f"  Forward Monthly Amount:          {money(fwd_monthly)}/month")
    doc.add_paragraph("")

    # ── Correction #6: 30-day cure period ──
    p = doc.add_paragraph()
    run = p.add_run("DEMAND FOR CURE:")
    run.bold = True

    # Refinement #6: Bold the cure deadline date within the paragraph
    if past_due is not None and past_due > 0:
        p = doc.add_paragraph()
        p.add_run(
            "You are hereby notified that you have thirty (30) days from receipt "
            "of this notice (cure deadline: "
        )
        run = p.add_run(CURE_DATE)
        run.bold = True
        p.add_run(
            f") to cure this default by "
            f"remitting the full arrears balance of {money(past_due)} and establishing "
            f"forward payment at the rate of {money(fwd_monthly)} per month."
        )
    else:
        p = doc.add_paragraph()
        p.add_run(
            "You are hereby notified that you have thirty (30) days from receipt "
            "of this notice (cure deadline: "
        )
        run = p.add_run(CURE_DATE)
        run.bold = True
        p.add_run(
            ") to cure this default by "
            "contacting us to reconcile the outstanding arrears balance and "
            f"establishing forward payment at the rate of {money(fwd_monthly)} per month."
        )

    doc.add_paragraph("")
    doc.add_paragraph(
        "FAILURE TO CURE this default by the cure deadline will result in "
        "the following actions:"
    )
    doc.add_paragraph(
        "  1. Filing of a Notice of Lien against the property under the "
        f"Declaration (target date: {LIEN_DEADLINE})"
    )
    doc.add_paragraph(
        "  2. Referral to counsel for enforcement of all rights under the "
        "Declaration, including recovery of attorneys\u2019 fees and costs "
        "as provided therein"
    )
    doc.add_paragraph(
        "  3. Notice to senior lenders and title companies of the recorded lien"
    )
    # Refinement #3: Fourth failure-to-cure item
    doc.add_paragraph(
        "  4. Notification to regulatory agencies, insurance carriers, and "
        "business partners of the security impairment at the property, as applicable."
    )
    doc.add_paragraph("")

    # Legal
    doc.add_paragraph(
        "This matter is governed by Tennessee law. The Declaration provides for "
        "lien rights, fee-shifting, and forum selection in Shelby County, Tennessee. "
        "The applicable statute of limitations for contract enforcement is six (6) years."
    )
    doc.add_paragraph("")
    # Refinement #4: Resolution conference language
    doc.add_paragraph(
        "We are prepared to discuss a reasonable resolution of the arrears balance, "
        "including structured payment arrangements, provided forward compliance is "
        "established immediately. Please contact this office within ten (10) business "
        "days of receipt to schedule a resolution conference."
    )
    doc.add_paragraph("")

    # ── Correction #8: Instrument citation above signature ──
    doc.add_paragraph(INSTRUMENT_CITATION)
    doc.add_paragraph("")

    # ── Correction #2: New signature block ──
    doc.add_paragraph("Respectfully,")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("____________________________________")
    doc.add_paragraph("Walter D. Wills III")
    doc.add_paragraph("Managing Partner, Wills & Wills LP")
    doc.add_paragraph("Declarant")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Prepared by Vanguard Security Services, Designated Agent per "
        "Security Management Agreement dated February 12, 2026."
    )
    doc.add_paragraph("")

    # ── Correction #3: Updated cc line ──
    doc.add_paragraph("cc: Jeff Rosenblum, Esq., Legal Counsel")

    return doc


def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    rows = conn.execute(
        "SELECT * FROM parcels WHERE status = 'DELINQUENT' ORDER BY sqft DESC"
    ).fetchall()

    print(f"  Generating demand letters for {len(rows)} delinquent targets (from live DB)...")
    print()
    print(f"  {'#':>2}  {'Target':<35} {'SqFt':>8} {'% Campus':>9} "
          f"{'Past Due':>14} {'$/Week':>10}  File")
    print("  " + "-" * 120)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    datestamp = datetime.now().strftime("%Y%m%d")

    for i, row in enumerate(rows, 1):
        sqft = row["sqft"] or 0
        pct = sqft / TOTAL_CAMPUS_SQFT
        past_due = row["past_due_balance"]
        weekly = row["weekly_rate"] or 0
        pd_str = money(past_due) if past_due and past_due > 0 else "TBD"

        doc = build_letter(row)

        safe_name = (row["business_name"]
                     .replace("/", "-").replace(" ", "_")
                     .replace("'", "").replace("\u2019", "")
                     .replace("#", "").replace("(", "").replace(")", ""))
        filename = f"Demand_{safe_name}_{datestamp}.docx"
        filepath = os.path.join(BASEDIR, filename)
        doc.save(filepath)

        # Log in enforcement_log
        conn.execute(
            """INSERT INTO enforcement_log
               (parcel_id, timestamp, action, next_step, attorney, notes)
               VALUES (?,?,?,?,?,?)""",
            (row["id"], now, f"Demand letter generated (v4-verified): {filename}",
             "Send via certified mail", "Rosenblum",
             f"Arrears: {pd_str}, SF: {sqft:,}, 30-day cure"),
        )

        print(f"  {i:>2}  {row['business_name']:<35} {sqft:>8,} {pct:>8.4%} "
              f"{pd_str:>14} {money(weekly):>10}  {filename}")

    conn.commit()

    total_known = sum(
        r["past_due_balance"] for r in rows
        if r["past_due_balance"] is not None and r["past_due_balance"] > 0
    )
    total_weekly = sum(r["weekly_rate"] or 0 for r in rows)
    print()
    print(f"  Total known arrears:   {money(total_known)}")
    print(f"  Total weekly rate:     {money(total_weekly)}")
    print(f"  Letters generated:     {len(rows)}")
    print(f"  Cure period:           30 days (deadline: {CURE_DATE})")
    print(f"  Saved to:              {BASEDIR}")

    conn.close()


if __name__ == "__main__":
    main()
