#!/usr/bin/env python3
"""
KIRBY GATE ENFORCEMENT SYSTEM
Covenant enforcement tracker for 21-parcel, 672,718 SF commercial campus.
All actions are timestamped and auditable. Runs entirely local on SQLite.
"""

import csv
import os
import sqlite3
import sys
import textwrap
from datetime import datetime, timedelta

# Force UTF-8 output on Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "kirbygate.db")

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════

TOTAL_CAMPUS_SQFT = 672_718
HISTORIC_WEEKLY_RATE = 6_069.52   # pre-Jan 2026
CURRENT_WEEKLY_RATE  = 9_000.00   # Jan 2026+
ARREARS_WEEKS = 156               # 36 months
LIEN_DEADLINE = "2026-04-01"
CURE_PERIOD_DAYS = 15
DECLARATION_DATE = "2011-05-05"

VALID_STATUSES = ["CURRENT", "DELINQUENT", "DISPUTED", "RECON", "VERIFY", "SETTLED"]
VALID_STEPS = [
    "Paying", "Demand Drafted", "Demand Sent", "Response Received",
    "In Negotiation", "Settlement Agreed", "Lien Filed",
    "Attorney Letter Received", "Research", "Resolved"
]

# ═══════════════════════════════════════════════════════════════════════════════
# DATABASE SETUP
# ═══════════════════════════════════════════════════════════════════════════════

SCHEMA = """
CREATE TABLE IF NOT EXISTS parcels (
    id                INTEGER PRIMARY KEY AUTOINCREMENT,
    address           TEXT NOT NULL,
    business_name     TEXT,
    sqft              INTEGER,
    pct_campus        REAL,
    status            TEXT DEFAULT 'VERIFY',
    entity_owner      TEXT,
    corporate_target  TEXT,
    past_due_balance  REAL,
    weekly_rate       REAL,
    certified_mail_tracking TEXT,
    date_packet_sent  TEXT,
    cure_deadline     TEXT,
    lien_filing_date  TEXT,
    attorney_referral_date TEXT,
    enforcement_step  TEXT DEFAULT 'Research',
    next_action       TEXT,
    deadline          TEXT,
    notes             TEXT,
    county_parcel_id  TEXT,
    mailing_address   TEXT,
    lender_name       TEXT,
    lender_address    TEXT,
    deed_of_trust_ref TEXT,
    lender_contact    TEXT,
    loan_number       TEXT,
    title_company     TEXT,
    address_verified  INTEGER DEFAULT 0,
    lender_verified   INTEGER DEFAULT 0
);

CREATE TABLE IF NOT EXISTS enforcement_log (
    id                INTEGER PRIMARY KEY AUTOINCREMENT,
    parcel_id         INTEGER,
    timestamp         TEXT NOT NULL,
    action            TEXT,
    sent_via          TEXT,
    response_due      TEXT,
    response_received TEXT,
    next_step         TEXT,
    attorney          TEXT,
    cost              REAL DEFAULT 0,
    notes             TEXT,
    FOREIGN KEY (parcel_id) REFERENCES parcels(id)
);

CREATE TABLE IF NOT EXISTS rates (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    label           TEXT,
    value           REAL,
    effective_date  TEXT
);
"""

SEED_PARCELS = [
    # CURRENT (paying)
    ("6500 Kirby Gate Blvd", "Waters of Memphis (Apts)", 54424, "CURRENT", "Kirby Gate Apts LLC", "", "Paying", "", "", ""),
    ("2809 Kirby Pkwy", "Shoppes at Kirby", 22350, "CURRENT", "WFC", "", "Paying", "", "", ""),
    ("2857 Kirby Pkwy", "Shops at Kirby Gate", 20000, "CURRENT", "WFC", "", "Paying", "", "", ""),
    ("2865 Kirby Pkwy", "Kirby Wines & Spirits", 5500, "CURRENT", "WFC", "", "Paying", "", "", ""),
    ("2873 Kirby Pkwy", "Nail Salon / Retail", 4200, "CURRENT", "WFC", "", "Paying", "", "", ""),
    ("6535 Kirby Gate Blvd", "Summit Medical (Bldg A)", 72000, "CURRENT", "Summit Healthcare REIT", "", "Paying", "", "", "Verify: may be partial"),
    ("6655 Quince Rd", "Willow Grove Apts", 72400, "CURRENT", "Willow Grove LP", "", "Paying", "", "", "VERIFY: listed current in case summary"),
    ("6660 Quince Rd", "Kroger #445", 65000, "CURRENT", "Kroger Co.", "Cincinnati HQ", "Paying", "", "", ""),
    ("2725 Kirby Pkwy", "FedEx Office / Retail", 18000, "CURRENT", "Kirby Pkwy Retail LLC", "", "Paying", "", "", ""),
    # DELINQUENT (non-paying)
    ("6480 Quince Rd", "Pointe at Kirby (MedHCP)", 31061, "DELINQUENT", "Grace Mgmt / MedHCP REIT", "MedHCP Corp - Dallas TX", "Demand Sent", "Follow-up / Lien Warning", "2026-03-01", "Large parcel, high-value target"),
    ("6500 Quince Rd", "Freedom Plasma Center", 12314, "DELINQUENT", "Realty Income Corp (NYSE: O)", "San Diego HQ - Corp Real Estate", "Demand Drafted", "Send certified demand", "2026-02-28", "Public REIT - will settle to avoid headline"),
    ("6532 Kirby Gate Blvd", "Dollar Tree #12847", 9964, "DELINQUENT", "Dollar Tree Inc.", "Chesapeake VA - Corp Real Estate", "Demand Drafted", "Send certified demand", "2026-02-28", "Corporate tenant, standard collections"),
    ("6659 Quince Rd", "Dollar General #22748", 9301, "DELINQUENT", "Dollar General Corp.", "Goodlettsville TN HQ", "Demand Drafted", "Send certified demand", "2026-02-28", "Corporate tenant, standard collections"),
    ("2801 Kirby Pkwy", "Starbucks (Exline prop)", 7493, "DELINQUENT", "Exline Properties", "Starbucks Regional Facilities", "Demand Drafted", "Send certified demand", "2026-02-28", "Landlord = Exline; Starbucks = subtenant"),
    ("2845 Kirby Pkwy", "Wendy's (Carlisle/Wendelta)", 3284, "DELINQUENT", "Wendelta Inc. / Carlisle Corp", "Paul Volpe CFO / Greg Jones COO", "Demand Drafted", "Send certified demand", "2026-02-28", "Franchisee entity"),
    ("2735 Kirby Pkwy", "Dunkin' (JP Foods LLC)", 2375, "DELINQUENT", "JP Foods LLC", "Peter Garner, 999 S Shady Grove", "Demand Drafted", "Send certified demand", "2026-02-28", "Local franchisee - smallest parcel"),
    # RECON
    ("2715 Kirby Pkwy", "KG Business Center", 43200, "RECON", "Unknown - research needed", "TBD", "Research", "Identify entity & send demand", "2026-02-28", "Non-payer, arrears TBD - need Wills data"),
    # DISPUTED
    ("6635 Quince Rd", "GALR Properties", 28500, "DISPUTED", "GALR LP", "Price (attorney) - dispute active", "Attorney Letter Received", "Rosenblum response", "2026-03-15", "Price letter opened 'service quality' door"),
    # VERIFY
    ("2966 Kirby Rd", "Car Wash (GX18)", 5000, "VERIFY", "Unknown", "TBD", "Research", "Verify if on Vanguard rolls", "2026-03-01", "Flagged in Wills spreadsheet - may not be covenant parcel"),
]

SEED_RATES = [
    ("Historic Campus Weekly Rate", HISTORIC_WEEKLY_RATE, "2022-12-01"),
    ("Current Campus Weekly Rate", CURRENT_WEEKLY_RATE, "2026-01-01"),
    ("Total Campus SqFt", TOTAL_CAMPUS_SQFT, "2022-12-01"),
    ("Arrears Period (weeks)", ARREARS_WEEKS, "2022-12-01"),
    ("Cure Period (days)", CURE_PERIOD_DAYS, "2011-05-05"),
]

SEED_LOG = [
    (None, "Enforcement tracker created — Vanguard/ARS system operational", "", "", "", "Finalize demand letters", "Rosenblum", 0, ""),
    (None, "TARGET: All demands sent via certified mail", "USPS Certified", "2026-03-15", "", "Wait 15 days for response", "Rosenblum", 0, "Declaration requires 15-day cure period"),
]


def ensure_deadline_columns(conn):
    """Add deadline-tracking columns to existing databases that lack them."""
    existing = {row[1] for row in conn.execute("PRAGMA table_info(parcels)").fetchall()}
    new_cols = [
        ("certified_mail_tracking", "TEXT"),
        ("date_packet_sent", "TEXT"),
        ("cure_deadline", "TEXT"),
        ("lien_filing_date", "TEXT"),
        ("attorney_referral_date", "TEXT"),
        ("county_parcel_id", "TEXT"),
        ("mailing_address", "TEXT"),
        ("lender_name", "TEXT"),
        ("lender_address", "TEXT"),
        ("deed_of_trust_ref", "TEXT"),
        ("lender_contact", "TEXT"),
        ("loan_number", "TEXT"),
        ("title_company", "TEXT"),
        ("address_verified", "INTEGER DEFAULT 0"),
        ("lender_verified", "INTEGER DEFAULT 0"),
    ]
    for col_name, col_type in new_cols:
        if col_name not in existing:
            conn.execute(f"ALTER TABLE parcels ADD COLUMN {col_name} {col_type}")
    conn.commit()


def get_db():
    """Open and optionally initialize the database."""
    fresh = not os.path.exists(DB_PATH)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    for stmt in SCHEMA.split(";"):
        stmt = stmt.strip()
        if stmt:
            conn.execute(stmt)
    conn.commit()

    # Migrate existing DBs to add deadline columns
    if not fresh:
        ensure_deadline_columns(conn)

    if fresh:
        # Seed parcels
        for p in SEED_PARCELS:
            sqft = p[2]
            pct = sqft / TOTAL_CAMPUS_SQFT
            conn.execute(
                """INSERT INTO parcels
                   (address, business_name, sqft, pct_campus, status,
                    entity_owner, corporate_target, enforcement_step,
                    next_action, deadline, notes)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                (p[0], p[1], sqft, pct, p[3], p[4], p[5], p[6], p[7], p[8], p[9]),
            )
        # Seed rates
        for r in SEED_RATES:
            conn.execute(
                "INSERT INTO rates (label, value, effective_date) VALUES (?,?,?)", r
            )
        # Seed enforcement log
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        for entry in SEED_LOG:
            conn.execute(
                """INSERT INTO enforcement_log
                   (parcel_id, timestamp, action, sent_via, response_due,
                    response_received, next_step, attorney, cost, notes)
                   VALUES (?,?,?,?,?,?,?,?,?,?)""",
                (entry[0], now, entry[1], entry[2], entry[3],
                 entry[4], entry[5], entry[6], entry[7], entry[8]),
            )
        conn.commit()
    return conn


# ═══════════════════════════════════════════════════════════════════════════════
# CALCULATION ENGINE
# ═══════════════════════════════════════════════════════════════════════════════

def calc_pct(sqft):
    return sqft / TOTAL_CAMPUS_SQFT

def calc_historic_weekly(sqft):
    return HISTORIC_WEEKLY_RATE * calc_pct(sqft)

def calc_current_weekly(sqft):
    return CURRENT_WEEKLY_RATE * calc_pct(sqft)

def calc_arrears_36mo(sqft):
    return calc_historic_weekly(sqft) * ARREARS_WEEKS

def calc_forward_monthly(sqft):
    return calc_current_weekly(sqft) * 4.333

def get_arrears(row):
    """Return stored past_due_balance if set, otherwise calculate from sqft."""
    if row["past_due_balance"] is not None and row["past_due_balance"] != 0:
        return row["past_due_balance"]
    if row["status"] in ("DELINQUENT", "DISPUTED", "RECON"):
        return calc_arrears_36mo(row["sqft"] or 0)
    return 0

def get_weekly(row):
    """Return stored weekly_rate if set, otherwise calculate from sqft."""
    if row["weekly_rate"] is not None and row["weekly_rate"] != 0:
        return row["weekly_rate"]
    return calc_current_weekly(row["sqft"] or 0)

def get_forward_monthly(row):
    """Forward monthly from stored weekly rate or calculated."""
    return get_weekly(row) * 4.333

def calc_deadlines(date_sent_str):
    """Given a packet sent date (YYYY-MM-DD), return cure/lien/attorney dates."""
    sent = datetime.strptime(date_sent_str, "%Y-%m-%d")
    return {
        "cure_deadline": (sent + timedelta(days=30)).strftime("%Y-%m-%d"),
        "lien_filing_date": (sent + timedelta(days=45)).strftime("%Y-%m-%d"),
        "attorney_referral_date": (sent + timedelta(days=60)).strftime("%Y-%m-%d"),
    }

def set_packet_sent(conn, parcel_id, date_sent_str, tracking_number=None):
    """Mark a parcel as packet sent and auto-calculate all deadlines."""
    deadlines = calc_deadlines(date_sent_str)
    conn.execute(
        """UPDATE parcels SET
           date_packet_sent = ?,
           certified_mail_tracking = COALESCE(?, certified_mail_tracking),
           cure_deadline = ?,
           lien_filing_date = ?,
           attorney_referral_date = ?,
           enforcement_step = 'Demand Sent',
           next_action = 'Await cure by ' || ?
           WHERE id = ?""",
        (date_sent_str, tracking_number,
         deadlines["cure_deadline"],
         deadlines["lien_filing_date"],
         deadlines["attorney_referral_date"],
         deadlines["cure_deadline"],
         parcel_id),
    )
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    tracking_note = f", tracking: {tracking_number}" if tracking_number else ""
    conn.execute(
        """INSERT INTO enforcement_log
           (parcel_id, timestamp, action, sent_via, response_due, next_step, attorney, notes)
           VALUES (?,?,?,?,?,?,?,?)""",
        (parcel_id, now,
         f"Demand packet sent on {date_sent_str}{tracking_note}",
         "USPS Certified",
         deadlines["cure_deadline"],
         f"Cure by {deadlines['cure_deadline']}, lien by {deadlines['lien_filing_date']}, attorney by {deadlines['attorney_referral_date']}",
         "Rosenblum",
         f"30-day cure: {deadlines['cure_deadline']} | 45-day lien: {deadlines['lien_filing_date']} | 60-day attorney: {deadlines['attorney_referral_date']}"),
    )
    conn.commit()
    return deadlines

def calc_settlement(principal, discount_pct, interest_rate, term_months):
    settled = principal * (1 - discount_pct)
    monthly_no_int = settled / term_months if term_months else 0
    total_with_int = settled * (1 + interest_rate * (term_months / 12))
    monthly_with_int = total_with_int / term_months if term_months else 0
    litigation_est = principal * 1.40  # 40% premium
    savings_vs_full = principal - settled
    savings_vs_lit = litigation_est - total_with_int
    return {
        "settled_amount": settled,
        "monthly_no_interest": monthly_no_int,
        "total_with_interest": total_with_int,
        "monthly_with_interest": monthly_with_int,
        "savings_vs_full": savings_vs_full,
        "savings_pct": discount_pct,
        "savings_vs_litigation": savings_vs_lit,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# DISPLAY HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def money(val):
    if val is None or val == 0:
        return "—"
    return f"${val:,.2f}"

def trunc(text, width):
    if text is None or text == "":
        return "—"
    text = str(text)
    return text if len(text) <= width else text[: width - 1] + "~"

def status_marker(status):
    markers = {
        "CURRENT": "[OK]",
        "DELINQUENT": "[!!]",
        "DISPUTED": "[??]",
        "RECON": "[RR]",
        "VERIFY": "[VV]",
        "SETTLED": "[$$]",
    }
    return markers.get(status, "[  ]")

def clear():
    os.system("cls" if os.name == "nt" else "clear")

def pause():
    input("\n  Press Enter to continue...")

def header(title):
    print()
    print("  " + "=" * 70)
    print(f"  KIRBY GATE ENFORCEMENT SYSTEM — {title}")
    print("  " + "=" * 70)
    print()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 1: VIEW ALL PARCELS
# ═══════════════════════════════════════════════════════════════════════════════

def view_all_parcels(conn):
    header("ALL PARCELS (21)")
    rows = conn.execute("SELECT * FROM parcels ORDER BY status, id").fetchall()

    hdr = (
        f"{'ID':>3} {'Stat':>5} {'Address':<24} {'Business':<28} "
        f"{'SqFt':>7} {'%':>6} {'36-Mo Arrears':>14} {'Wkly Now':>10} "
        f"{'Step':<16} {'Deadline':<11}"
    )
    sep = "  " + "-" * len(hdr)
    print(f"  {hdr}")
    print(sep)

    total_sqft = 0
    total_arrears = 0
    for r in rows:
        sqft = r["sqft"] or 0
        pct = r["pct_campus"] or 0
        arrears = get_arrears(r)
        wkly = get_weekly(r)
        total_sqft += sqft
        total_arrears += arrears

        line = (
            f"{r['id']:>3} {status_marker(r['status']):>5} "
            f"{trunc(r['address'], 24):<24} {trunc(r['business_name'], 28):<28} "
            f"{sqft:>7,} {pct:>5.1%} {money(arrears):>14} {money(wkly):>10} "
            f"{trunc(r['enforcement_step'], 16):<16} {trunc(r['deadline'], 11):<11}"
        )
        print(f"  {line}")

    print(sep)
    print(f"  {'':>3} {'':>5} {'TOTALS':<24} {'':28} {total_sqft:>7,} "
          f"{'':>6} {money(total_arrears):>14}")
    print()
    print(f"  Legend:  [OK] Current   [!!] Delinquent   [??] Disputed   [RR] Recon   [VV] Verify   [$$] Settled")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 2: VIEW NON-PAYERS ONLY
# ═══════════════════════════════════════════════════════════════════════════════

def view_nonpayers(conn):
    header("NON-PAYERS (Delinquent / Disputed / Recon)")
    rows = conn.execute(
        "SELECT * FROM parcels WHERE status IN ('DELINQUENT','DISPUTED','RECON') ORDER BY sqft DESC"
    ).fetchall()

    if not rows:
        print("  No non-paying parcels found.")
        pause()
        return

    total_arrears = 0
    total_weekly = 0
    total_monthly = 0
    print(f"  {'ID':>3} {'Address':<24} {'Business':<28} {'SqFt':>7} "
          f"{'36-Mo Arrears':>14} {'Wkly Now':>10} {'Fwd Mo':>10} "
          f"{'Corporate Target':<30}")
    print("  " + "-" * 130)

    for r in rows:
        sqft = r["sqft"] or 0
        arrears = get_arrears(r)
        wkly = get_weekly(r)
        mo = get_forward_monthly(r)
        total_arrears += arrears
        total_weekly += wkly
        total_monthly += mo

        print(f"  {r['id']:>3} {trunc(r['address'], 24):<24} "
              f"{trunc(r['business_name'], 28):<28} {sqft:>7,} "
              f"{money(arrears):>14} {money(wkly):>10} {money(mo):>10} "
              f"{trunc(r['corporate_target'], 30):<30}")

    print("  " + "-" * 130)
    print(f"  {'':>3} {'TOTALS':<24} {'':28} {'':>7} "
          f"{money(total_arrears):>14} {money(total_weekly):>10} {money(total_monthly):>10}")
    print()
    print(f"  Non-payer count: {len(rows)}")
    print(f"  Lien deadline:   {LIEN_DEADLINE}")
    days_left = (datetime.strptime(LIEN_DEADLINE, "%Y-%m-%d") - datetime.now()).days
    print(f"  Days remaining:  {days_left}")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 3: UPDATE PARCEL STATUS
# ═══════════════════════════════════════════════════════════════════════════════

def update_parcel(conn):
    header("UPDATE PARCEL")
    pid = input("  Enter parcel ID (or 'list' to see IDs): ").strip()
    if pid.lower() == "list":
        rows = conn.execute("SELECT id, address, business_name FROM parcels ORDER BY id").fetchall()
        for r in rows:
            print(f"    {r['id']:>3}  {r['address']:<26} {r['business_name']}")
        print()
        pid = input("  Enter parcel ID: ").strip()

    try:
        pid = int(pid)
    except ValueError:
        print("  Invalid ID.")
        pause()
        return

    row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
    if not row:
        print(f"  No parcel with ID {pid}.")
        pause()
        return

    # Show current values
    print(f"\n  Parcel #{row['id']}: {row['business_name']}")
    print(f"  Address:          {row['address']}")
    print(f"  Status:           {row['status']}")
    print(f"  Enforcement Step: {row['enforcement_step']}")
    print(f"  Next Action:      {row['next_action'] or '—'}")
    print(f"  Deadline:         {row['deadline'] or '—'}")
    print(f"  Entity/Owner:     {row['entity_owner'] or '—'}")
    print(f"  Corporate Target: {row['corporate_target'] or '—'}")
    print(f"  Notes:            {row['notes'] or '—'}")

    print("\n  What do you want to update?")
    print("    1. Status")
    print("    2. Enforcement Step")
    print("    3. Next Action")
    print("    4. Deadline")
    print("    5. Corporate Target")
    print("    6. Entity/Owner")
    print("    7. Notes")
    print("    8. Business Name")
    print("    9. Square Footage")
    print("   10. Past Due Balance")
    print("   11. Weekly Rate")
    print("  ---- Lender / Title ----")
    print("   12. County Parcel ID")
    print("   13. Mailing Address (for certified mail)")
    print("   14. Lender Name")
    print("   15. Lender Address")
    print("   16. Deed of Trust Reference (book/page)")
    print("   17. Lender Contact")
    print("   18. Loan Number")
    print("   19. Title Company")
    print("    0. Cancel")

    choice = input("\n  Choice: ").strip()

    field_map = {
        "1": ("status", "Status", VALID_STATUSES),
        "2": ("enforcement_step", "Enforcement Step", VALID_STEPS),
        "3": ("next_action", "Next Action", None),
        "4": ("deadline", "Deadline (YYYY-MM-DD)", None),
        "5": ("corporate_target", "Corporate Target", None),
        "6": ("entity_owner", "Entity/Owner", None),
        "7": ("notes", "Notes", None),
        "8": ("business_name", "Business Name", None),
        "9": ("sqft", "Square Footage", None),
        "10": ("past_due_balance", "Past Due Balance", None),
        "11": ("weekly_rate", "Weekly Rate", None),
        "12": ("county_parcel_id", "County Parcel ID", None),
        "13": ("mailing_address", "Mailing Address", None),
        "14": ("lender_name", "Lender Name", None),
        "15": ("lender_address", "Lender Address", None),
        "16": ("deed_of_trust_ref", "Deed of Trust Ref (book/page)", None),
        "17": ("lender_contact", "Lender Contact", None),
        "18": ("loan_number", "Loan Number", None),
        "19": ("title_company", "Title Company", None),
    }

    if choice == "0" or choice not in field_map:
        return

    field, label, options = field_map[choice]

    if options:
        print(f"\n  Valid values for {label}:")
        for i, opt in enumerate(options, 1):
            print(f"    {i}. {opt}")
        sel = input(f"\n  Select {label} (number or type value): ").strip()
        try:
            idx = int(sel) - 1
            if 0 <= idx < len(options):
                value = options[idx]
            else:
                value = sel
        except ValueError:
            value = sel
    else:
        value = input(f"\n  Enter new {label}: ").strip()

    # Handle numeric fields
    if field == "sqft":
        try:
            value = int(value.replace(",", ""))
        except ValueError:
            print("  Invalid number.")
            pause()
            return
        pct = value / TOTAL_CAMPUS_SQFT
        conn.execute("UPDATE parcels SET sqft = ?, pct_campus = ? WHERE id = ?", (value, pct, pid))
    elif field in ("past_due_balance", "weekly_rate"):
        try:
            value = float(value.replace("$", "").replace(",", ""))
        except ValueError:
            print("  Invalid number.")
            pause()
            return
        conn.execute(f"UPDATE parcels SET {field} = ? WHERE id = ?", (value, pid))
    else:
        conn.execute(f"UPDATE parcels SET {field} = ? WHERE id = ?", (value, pid))

    # Auto-log the change
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        """INSERT INTO enforcement_log (parcel_id, timestamp, action, notes)
           VALUES (?, ?, ?, ?)""",
        (pid, now, f"Updated {label} to: {value}", f"Changed by user"),
    )
    conn.commit()
    print(f"\n  Updated {row['business_name']} — {label} = {value}")
    print(f"  (Logged at {now})")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 4: LOG ENFORCEMENT ACTION
# ═══════════════════════════════════════════════════════════════════════════════

def log_enforcement(conn):
    header("LOG ENFORCEMENT ACTION")
    pid = input("  Parcel ID (or 'all' for campus-wide, 'list' to see IDs): ").strip()

    if pid.lower() == "list":
        rows = conn.execute("SELECT id, address, business_name FROM parcels ORDER BY id").fetchall()
        for r in rows:
            print(f"    {r['id']:>3}  {r['address']:<26} {r['business_name']}")
        print()
        pid = input("  Parcel ID (or 'all'): ").strip()

    if pid.lower() == "all":
        parcel_id = None
        parcel_name = "ALL PARCELS"
    else:
        try:
            parcel_id = int(pid)
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
        row = conn.execute("SELECT business_name FROM parcels WHERE id = ?", (parcel_id,)).fetchone()
        if not row:
            print(f"  No parcel with ID {parcel_id}.")
            pause()
            return
        parcel_name = row["business_name"]

    print(f"\n  Logging action for: {parcel_name}")
    action = input("  Action taken: ").strip()
    if not action:
        print("  Cancelled — no action entered.")
        pause()
        return

    print("\n  Sent via:")
    print("    1. USPS Certified Mail")
    print("    2. Email")
    print("    3. Hand-delivered")
    print("    4. FedEx / UPS")
    print("    5. Attorney (Rosenblum)")
    print("    6. N/A")
    sent_choice = input("  Choice: ").strip()
    sent_map = {"1": "USPS Certified", "2": "Email", "3": "Hand-delivered",
                "4": "FedEx/UPS", "5": "Attorney (Rosenblum)", "6": ""}
    sent_via = sent_map.get(sent_choice, sent_choice)

    response_due = input("  Response due date (YYYY-MM-DD, or Enter to skip): ").strip()
    next_step = input("  Next step: ").strip()
    attorney = input("  Attorney (Enter for Rosenblum): ").strip() or "Rosenblum"
    cost_str = input("  Cost ($ amount, or Enter for 0): ").strip()
    try:
        cost = float(cost_str.replace("$", "").replace(",", "")) if cost_str else 0
    except ValueError:
        cost = 0
    notes = input("  Notes: ").strip()

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        """INSERT INTO enforcement_log
           (parcel_id, timestamp, action, sent_via, response_due,
            response_received, next_step, attorney, cost, notes)
           VALUES (?,?,?,?,?,?,?,?,?,?)""",
        (parcel_id, now, action, sent_via, response_due or None,
         None, next_step, attorney, cost, notes),
    )
    conn.commit()
    print(f"\n  Action logged at {now}")
    print(f"  Parcel:   {parcel_name}")
    print(f"  Action:   {action}")
    print(f"  Sent via: {sent_via or 'N/A'}")
    if response_due:
        print(f"  Response due: {response_due}")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 5: GENERATE DEMAND LETTER
# ═══════════════════════════════════════════════════════════════════════════════

def generate_demand_letter(conn):
    header("GENERATE DEMAND LETTER")

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  Error: python-docx is not installed.")
        print("  Run: pip install python-docx")
        pause()
        return

    pid = input("  Parcel ID (or 'list' to see non-payers): ").strip()
    if pid.lower() == "list":
        rows = conn.execute(
            "SELECT id, address, business_name FROM parcels WHERE status != 'CURRENT' ORDER BY id"
        ).fetchall()
        for r in rows:
            print(f"    {r['id']:>3}  {r['address']:<26} {r['business_name']}")
        print()
        pid = input("  Parcel ID: ").strip()

    try:
        pid = int(pid)
    except ValueError:
        print("  Invalid ID.")
        pause()
        return

    row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
    if not row:
        print(f"  No parcel with ID {pid}.")
        pause()
        return

    sqft = row["sqft"] or 0
    pct = row["pct_campus"] or 0
    arrears = get_arrears(row)
    curr_wkly = get_weekly(row)
    fwd_monthly = get_forward_monthly(row)
    today = datetime.now().strftime("%B %d, %Y")
    cure_date = (datetime.now() + timedelta(days=CURE_PERIOD_DAYS)).strftime("%B %d, %Y")

    doc = Document()

    # Letterhead placeholder
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VANGUARD SECURITY SERVICES / ARS")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[Letterhead — Address / Phone / Email]")

    doc.add_paragraph("")
    doc.add_paragraph(today)
    doc.add_paragraph("")

    # Addressee
    if row["corporate_target"]:
        doc.add_paragraph(f"SENT VIA CERTIFIED MAIL")
        doc.add_paragraph("")
        doc.add_paragraph(f"{row['corporate_target']}")
    if row["entity_owner"]:
        doc.add_paragraph(f"RE: {row['entity_owner']}")
    doc.add_paragraph(f"Property: {row['address']}")
    doc.add_paragraph(f"Tenant/Occupant: {row['business_name']}")
    doc.add_paragraph("")

    # Subject
    p = doc.add_paragraph()
    run = p.add_run("RE: NOTICE OF COVENANT NON-COMPLIANCE AND DEMAND FOR CURE")
    run.bold = True
    run.underline = True
    doc.add_paragraph("")

    # Body
    doc.add_paragraph(
        f"Dear Sir or Madam:"
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        f"This letter constitutes formal notice that the above-referenced property is in material "
        f"default of the Declaration of Restrictive Covenants recorded on {DECLARATION_DATE} "
        f"(the \"Declaration\") governing the Kirby Gate commercial development, Memphis, Tennessee."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Pursuant to the Declaration, each parcel owner is obligated to fund its pro-rata share of "
        f"campus-wide security services. Vanguard Security Services has been designated as the "
        f"authorized security provider under the Declaration, as confirmed by the Wills designation letter."
    )
    doc.add_paragraph("")

    # Calculation breakdown
    p = doc.add_paragraph()
    run = p.add_run("ARREARS CALCULATION:")
    run.bold = True

    doc.add_paragraph(f"  Total Campus Square Footage:   {TOTAL_CAMPUS_SQFT:,} SF")
    doc.add_paragraph(f"  Your Parcel Square Footage:    {sqft:,} SF")
    doc.add_paragraph(f"  Your Pro-Rata Share:           {pct:.4%}")
    doc.add_paragraph(f"  Your Weekly Rate:              {money(curr_wkly)}/week")
    doc.add_paragraph(f"  Arrears Period:                {ARREARS_WEEKS} weeks (36 months)")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    if arrears:
        run = p.add_run(f"  TOTAL 36-MONTH ARREARS OWED:   {money(arrears)}")
    else:
        run = p.add_run("  TOTAL 36-MONTH ARREARS OWED:   TO BE DETERMINED (pending reconciliation)")
    run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("FORWARD BILLING (Effective January 2026):")
    run.bold = True
    doc.add_paragraph(f"  Current Weekly Rate:           {money(curr_wkly)}/week")
    doc.add_paragraph(f"  Forward Monthly Amount:        {money(fwd_monthly)}/month")
    doc.add_paragraph("")

    # Demand
    p = doc.add_paragraph()
    run = p.add_run("DEMAND FOR CURE:")
    run.bold = True
    doc.add_paragraph(
        f"You are hereby notified that you have fifteen (15) days from receipt of this notice "
        f"(cure deadline: {cure_date}) to cure this default by remitting the full arrears "
        f"balance of {money(arrears)} and establishing forward payment at the rate of "
        f"{money(fwd_monthly)} per month."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        f"FAILURE TO CURE this default by the cure deadline will result in the following actions:"
    )
    doc.add_paragraph(
        f"  1. Filing of a Notice of Lien against the property under the Declaration (target date: {LIEN_DEADLINE})"
    )
    doc.add_paragraph(
        f"  2. Referral to counsel for enforcement of all rights under the Declaration, "
        f"including recovery of attorneys' fees and costs as provided therein"
    )
    doc.add_paragraph(
        f"  3. Notice to senior lenders and title companies of the recorded lien"
    )
    doc.add_paragraph("")

    doc.add_paragraph(
        f"This matter is governed by Tennessee law. The Declaration provides for lien rights, "
        f"fee-shifting, and forum selection in Shelby County, Tennessee. The applicable statute "
        f"of limitations for contract enforcement is six (6) years."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        f"We encourage you to contact us promptly to discuss resolution of this matter."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Respectfully,")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("____________________________________")
    doc.add_paragraph("Brad")
    doc.add_paragraph("Vanguard Security Services")
    doc.add_paragraph("")
    doc.add_paragraph("cc: Rosenblum (Counsel)")

    # Save
    safe_name = row["business_name"].replace("/", "-").replace(" ", "_")
    filename = f"Demand_{safe_name}_{datetime.now():%Y%m%d}.docx"
    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
    doc.save(filepath)

    # Log the action
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        """INSERT INTO enforcement_log
           (parcel_id, timestamp, action, next_step, attorney, notes)
           VALUES (?,?,?,?,?,?)""",
        (pid, now, f"Demand letter generated: {filename}",
         "Send via certified mail", "Rosenblum", f"Arrears: {money(arrears)}"),
    )
    conn.commit()

    print(f"\n  Demand letter saved: {filepath}")
    print(f"  Parcel:    {row['business_name']}")
    print(f"  Arrears:   {money(arrears)}")
    print(f"  Fwd Mo:    {money(fwd_monthly)}")
    print(f"  Cure date: {cure_date}")
    print(f"  (Action logged)")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 6: SETTLEMENT CALCULATOR
# ═══════════════════════════════════════════════════════════════════════════════

def settlement_calculator(conn):
    header("SETTLEMENT CALCULATOR")

    pid = input("  Parcel ID (or Enter to use custom amount): ").strip()
    if pid:
        try:
            pid = int(pid)
            row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
            if not row:
                print(f"  No parcel with ID {pid}.")
                pause()
                return
            principal = get_arrears(row)
            print(f"  Parcel: {row['business_name']}")
            print(f"  36-month arrears: {money(principal)}")
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
    else:
        amt = input("  Enter total arrears owed: $").strip().replace(",", "")
        try:
            principal = float(amt)
        except ValueError:
            print("  Invalid amount.")
            pause()
            return

    print()
    disc_str = input("  Discount offered (%, e.g. 35 for 35%): ").strip()
    try:
        discount_pct = float(disc_str) / 100
    except ValueError:
        discount_pct = 0.35
        print(f"  Using default: 35%")

    int_str = input("  Interest rate (%, e.g. 2 for 2%): ").strip()
    try:
        interest_rate = float(int_str) / 100
    except ValueError:
        interest_rate = 0.02
        print(f"  Using default: 2%")

    term_str = input("  Payment term (months, e.g. 36): ").strip()
    try:
        term = int(term_str)
    except ValueError:
        term = 36
        print(f"  Using default: 36 months")

    result = calc_settlement(principal, discount_pct, interest_rate, term)

    print()
    print("  " + "=" * 50)
    print("  SETTLEMENT TERMS")
    print("  " + "=" * 50)
    print(f"  Total Arrears Owed:        {money(principal)}")
    print(f"  Discount:                  {discount_pct:.0%}")
    print(f"  Settled Amount:            {money(result['settled_amount'])}")
    print(f"  Interest Rate:             {interest_rate:.1%}")
    print(f"  Term:                      {term} months")
    print("  " + "-" * 50)
    print(f"  Monthly Payment (no int):  {money(result['monthly_no_interest'])}")
    print(f"  Total w/ Interest:         {money(result['total_with_interest'])}")
    print(f"  Monthly Payment (w/ int):  {money(result['monthly_with_interest'])}")
    print("  " + "-" * 50)
    print(f"  Savings vs Full Arrears:   {money(result['savings_vs_full'])} ({result['savings_pct']:.0%})")
    print(f"  Savings vs Litigation*:    {money(result['savings_vs_litigation'])}")
    print()
    print("  * Litigation estimate assumes 40% cost premium (legal fees + time)")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 7: EXPORT TO EXCEL
# ═══════════════════════════════════════════════════════════════════════════════

def export_excel(conn):
    header("EXPORT TO EXCEL")

    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    except ImportError:
        print("  Error: openpyxl is not installed.")
        print("  Run: pip install openpyxl")
        pause()
        return

    wb = Workbook()
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(bold=True, size=11, color="FFFFFF")

    # Status fills
    fills = {
        "CURRENT": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
        "DELINQUENT": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
        "DISPUTED": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
        "RECON": PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid"),
        "VERIFY": PatternFill(start_color="BDD7EE", end_color="BDD7EE", fill_type="solid"),
        "SETTLED": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
    }

    # ── Sheet 1: Parcel Master ──
    ws1 = wb.active
    ws1.title = "Parcel Master"
    headers1 = [
        "Address", "Business Name", "SqFt", "% of Campus", "Status",
        "Historic Wkly Share", "36-Mo Arrears (156 wks)", "Current Wkly Share",
        "Fwd Monthly", "Entity/Owner", "Corporate Target", "Enforcement Step",
        "Next Action", "Deadline", "Notes"
    ]
    for col, h in enumerate(headers1, 1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.border = thin_border

    rows = conn.execute("SELECT * FROM parcels ORDER BY id").fetchall()
    for ri, r in enumerate(rows, 2):
        sqft = r["sqft"] or 0
        arrears = get_arrears(r)
        curr_wkly = get_weekly(r)
        fwd_mo = get_forward_monthly(r)
        hist_wkly = calc_historic_weekly(sqft)

        vals = [
            r["address"], r["business_name"], sqft, r["pct_campus"],
            r["status"], hist_wkly, arrears, curr_wkly, fwd_mo,
            r["entity_owner"], r["corporate_target"], r["enforcement_step"],
            r["next_action"], r["deadline"], r["notes"]
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws1.cell(row=ri, column=ci, value=v)
            cell.border = thin_border
            if r["status"] in fills:
                cell.fill = fills[r["status"]]

        # Format numbers
        ws1.cell(row=ri, column=4).number_format = '0.00%'
        for col_idx in (6, 7, 8, 9):
            ws1.cell(row=ri, column=col_idx).number_format = '$#,##0.00'

    # Autofit-ish widths
    widths = [24, 30, 8, 10, 14, 16, 18, 16, 14, 30, 32, 18, 28, 12, 40]
    for i, w in enumerate(widths, 1):
        ws1.column_dimensions[ws1.cell(row=1, column=i).column_letter].width = w

    # ── Sheet 2: Enforcement Timeline ──
    ws2 = wb.create_sheet("Enforcement Timeline")
    headers2 = [
        "Date", "Parcel Address", "Business Name", "Action Taken",
        "Sent Via", "Response Due", "Response Received", "Next Step",
        "Attorney", "Cost", "Notes"
    ]
    for col, h in enumerate(headers2, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.border = thin_border

    log_rows = conn.execute(
        """SELECT el.*, p.address, p.business_name
           FROM enforcement_log el
           LEFT JOIN parcels p ON el.parcel_id = p.id
           ORDER BY el.timestamp DESC"""
    ).fetchall()
    for ri, r in enumerate(log_rows, 2):
        vals = [
            r["timestamp"],
            r["address"] or "ALL PARCELS",
            r["business_name"] or "ALL",
            r["action"], r["sent_via"], r["response_due"],
            r["response_received"], r["next_step"],
            r["attorney"], r["cost"], r["notes"]
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws2.cell(row=ri, column=ci, value=v)
            cell.border = thin_border
        ws2.cell(row=ri, column=10).number_format = '$#,##0.00'

    widths2 = [20, 24, 28, 40, 16, 14, 14, 30, 14, 10, 40]
    for i, w in enumerate(widths2, 1):
        ws2.column_dimensions[ws2.cell(row=1, column=i).column_letter].width = w

    # ── Sheet 3: Settlement Calculator ──
    ws3 = wb.create_sheet("Settlement Calculator")
    ws3.cell(row=1, column=1, value="KIRBY GATE - SETTLEMENT CALCULATOR").font = Font(bold=True, size=14)
    labels = [
        (3, "INPUT PARAMETERS:"),
        (4, "Parcel Address"), (5, "Total Arrears Owed"), (6, "Discount Offered (%)"),
        (7, "Interest Rate (%)"), (8, "Payment Term (months)"),
        (10, "CALCULATED OUTPUTS:"),
        (11, "Settled Amount"), (12, "Monthly Payment (no interest)"),
        (13, "Total w/ Interest"), (14, "Monthly Payment (w/ interest)"),
        (15, "Savings vs Full Arrears"), (16, "Savings %"),
        (17, "Savings vs Litigation Est."),
        (19, "Note: Litigation est. assumes 40% premium (legal fees + time)"),
    ]
    for row_num, label in labels:
        ws3.cell(row=row_num, column=1, value=label)
    # Default values
    ws3.cell(row=5, column=2, value=50000)
    ws3.cell(row=6, column=2, value=0.35)
    ws3.cell(row=7, column=2, value=0.02)
    ws3.cell(row=8, column=2, value=36)
    ws3.column_dimensions["A"].width = 30
    ws3.column_dimensions["B"].width = 18

    # ── Sheet 4: Rates & Constants ──
    ws4 = wb.create_sheet("Rates & Constants")
    ws4.cell(row=1, column=1, value="KIRBY GATE - KEY RATES & CONSTANTS").font = Font(bold=True, size=14)
    constants = [
        (3, "Total Campus SqFt", TOTAL_CAMPUS_SQFT, "21 parcels total"),
        (4, "Historic Weekly Rate (campus)", HISTORIC_WEEKLY_RATE, "Pre-Jan 2026 rate"),
        (5, "Current Weekly Rate (campus)", CURRENT_WEEKLY_RATE, "Jan 2026+ rate"),
        (6, "Arrears Period (weeks)", ARREARS_WEEKS, "36 months x 4.333"),
        (7, "Arrears Period (months)", 36, "3-year lookback"),
        (8, "Statute of Limitations (years)", 6, "Tennessee contract SOL"),
        (9, "Max SOL Period (months)", 72, "Full 6-year lookback possible"),
        (10, "Lien Default Trigger (days)", CURE_PERIOD_DAYS, "Declaration cure period"),
        (11, "Lien Filing Deadline", LIEN_DEADLINE, "HARD DEADLINE"),
        (12, "Declaration Recorded", DECLARATION_DATE, "Instrument # needed from Rosenblum"),
        (13, "Paying Parcels", 9, "Current / compliant"),
        (14, "Non-Paying Parcels", 7, "Delinquent - demand needed"),
        (15, "Disputed Parcels", 1, "GALR - attorney involved"),
        (16, "Recon Needed", 2, "KG Business + Car Wash"),
        (17, "Attorney", "Rosenblum", "Lead counsel - enforcement"),
        (18, "Wills Designation", "Shannon Wills", "Original spreadsheet / data source"),
    ]
    for row_num, label, val, note in constants:
        ws4.cell(row=row_num, column=1, value=label)
        ws4.cell(row=row_num, column=2, value=val)
        ws4.cell(row=row_num, column=3, value=note)

    ws4.cell(row=21, column=1, value="COLOR LEGEND:").font = Font(bold=True)
    legend = [
        (22, "Green", "Current / Paying", "C6EFCE"),
        (23, "Red", "Delinquent - Demand Required", "FFC7CE"),
        (24, "Yellow", "Disputed - Attorney Involved", "FFEB9C"),
        (25, "Blue", "Recon Needed / Verify", "BDD7EE"),
    ]
    for row_num, color, desc, hex_color in legend:
        c1 = ws4.cell(row=row_num, column=1, value=color)
        c1.fill = PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")
        ws4.cell(row=row_num, column=2, value=desc)

    ws4.column_dimensions["A"].width = 30
    ws4.column_dimensions["B"].width = 18
    ws4.column_dimensions["C"].width = 40

    # Save
    filename = f"KirbyGate_Export_{datetime.now():%Y%m%d_%H%M%S}.xlsx"
    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
    wb.save(filepath)

    print(f"  Exported to: {filepath}")
    print(f"  Sheets: Parcel Master, Enforcement Timeline, Settlement Calculator, Rates & Constants")
    print(f"  Parcels: {len(rows)} | Log entries: {len(log_rows)}")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 8: ENFORCEMENT TIMELINE
# ═══════════════════════════════════════════════════════════════════════════════

def view_timeline(conn):
    header("ENFORCEMENT TIMELINE")
    rows = conn.execute(
        """SELECT el.*, p.address, p.business_name
           FROM enforcement_log el
           LEFT JOIN parcels p ON el.parcel_id = p.id
           ORDER BY el.timestamp DESC"""
    ).fetchall()

    if not rows:
        print("  No enforcement actions logged yet.")
        pause()
        return

    print(f"  {'Date':<20} {'Parcel':<26} {'Action':<40} {'Via':<16} {'Due':<12}")
    print("  " + "-" * 120)
    for r in rows:
        parcel = r["address"] or "ALL PARCELS"
        print(
            f"  {trunc(r['timestamp'], 20):<20} {trunc(parcel, 26):<26} "
            f"{trunc(r['action'], 40):<40} {trunc(r['sent_via'], 16):<16} "
            f"{trunc(r['response_due'], 12):<12}"
        )
        if r["next_step"]:
            print(f"  {'':20} {'':26} -> Next: {r['next_step']}")

    print()
    print(f"  Total actions logged: {len(rows)}")

    # Upcoming deadlines
    print()
    print("  KEY DEADLINES:")
    deadlines = conn.execute(
        """SELECT address, business_name, deadline, next_action
           FROM parcels WHERE deadline IS NOT NULL AND deadline != ''
           ORDER BY deadline"""
    ).fetchall()
    for d in deadlines:
        days = (datetime.strptime(d["deadline"], "%Y-%m-%d") - datetime.now()).days
        flag = " ** OVERDUE **" if days < 0 else f" ({days} days)"
        print(f"    {d['deadline']}  {d['address']:<24} {d['next_action'] or '—'}{flag}")

    print(f"\n    {LIEN_DEADLINE}  *** HARD DEADLINE: FILE LIENS ON NON-RESPONDERS ***")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 9: DASHBOARD SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

def dashboard(conn):
    header("DASHBOARD SUMMARY")

    all_rows = conn.execute("SELECT * FROM parcels").fetchall()
    current = [r for r in all_rows if r["status"] == "CURRENT"]
    delinquent = [r for r in all_rows if r["status"] == "DELINQUENT"]
    disputed = [r for r in all_rows if r["status"] == "DISPUTED"]
    recon = [r for r in all_rows if r["status"] in ("RECON", "VERIFY")]
    settled = [r for r in all_rows if r["status"] == "SETTLED"]

    total_campus_sqft = sum(r["sqft"] or 0 for r in all_rows)
    delinq_arrears = sum(get_arrears(r) for r in delinquent)
    disputed_arrears = sum(get_arrears(r) for r in disputed)
    recon_arrears = sum(get_arrears(r) for r in recon)
    total_arrears = delinq_arrears + disputed_arrears + recon_arrears
    delinq_weekly = sum(get_weekly(r) for r in delinquent)
    total_weekly = sum(get_weekly(r) for r in all_rows)

    days_to_lien = (datetime.strptime(LIEN_DEADLINE, "%Y-%m-%d") - datetime.now()).days

    print(f"  PARCELS:  {len(all_rows)} total  |  {len(current)} current  |  "
          f"{len(delinquent)} delinquent  |  {len(disputed)} disputed  |  "
          f"{len(recon)} recon/verify  |  {len(settled)} settled")
    print()
    print(f"  CAMPUS:   {total_campus_sqft:,} SF tracked of {TOTAL_CAMPUS_SQFT:,} SF total")
    print()
    print(f"  MONEY OWED:")
    print(f"    Delinquent arrears (36 mo):  {money(delinq_arrears)}")
    print(f"    Disputed arrears:            {money(disputed_arrears)}")
    print(f"    Recon/Verify arrears:        {money(recon_arrears)}")
    print(f"    ------------------------------------------")
    print(f"    TOTAL ARREARS:               {money(total_arrears)}")
    print()
    print(f"  WEEKLY REVENUE:")
    print(f"    Current weekly (campus):     {money(total_weekly)}")
    print(f"    Delinquent weekly unpaid:    {money(delinq_weekly)}")
    print()
    print(f"  LIEN DEADLINE:  {LIEN_DEADLINE}  ({days_to_lien} days remaining)")
    print()

    # Priority ranking by arrears
    print("  PRIORITY RANKING (by arrears amount):")
    print(f"  {'#':>3} {'Business':<30} {'Arrears':>14} {'Status':<12}")
    print("  " + "-" * 64)
    ranked = sorted(
        [r for r in all_rows if r["status"] != "CURRENT"],
        key=lambda r: get_arrears(r),
        reverse=True,
    )
    for i, r in enumerate(ranked, 1):
        arr = get_arrears(r)
        print(f"  {i:>3} {trunc(r['business_name'], 30):<30} {money(arr):>14} {r['status']:<12}")

    # Log count
    log_count = conn.execute("SELECT COUNT(*) as c FROM enforcement_log").fetchone()["c"]
    print(f"\n  AUDIT LOG: {log_count} enforcement actions recorded")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 10: PRO-RATA SECURITY CHARGE CALCULATOR
# ═══════════════════════════════════════════════════════════════════════════════

def prorata_calculator(conn):
    header("PRO-RATA SECURITY CHARGE CALCULATOR")

    print(f"  Declaration Formula:  (Parcel SF / {TOTAL_CAMPUS_SQFT:,} SF) x ${CURRENT_WEEKLY_RATE:,.2f}/week")
    print()

    rows = conn.execute(
        "SELECT * FROM parcels WHERE status = 'DELINQUENT' ORDER BY sqft DESC"
    ).fetchall()

    if not rows:
        print("  No delinquent parcels found.")
        pause()
        return

    # Table header
    print(f"  {'ID':>3}  {'Business':<30} {'Address':<22} "
          f"{'SqFt':>8} {'% Campus':>9} {'Pro-Rata Wk':>12} {'Billed Wk':>11} "
          f"{'Diff/Wk':>10} {'Pro-Rata Mo':>12} {'Billed Mo':>11}")
    print("  " + "-" * 155)

    total_prorata_wk = 0
    total_billed_wk = 0
    total_prorata_mo = 0
    total_billed_mo = 0

    for r in rows:
        sqft = r["sqft"] or 0
        pct = sqft / TOTAL_CAMPUS_SQFT
        prorata_wk = CURRENT_WEEKLY_RATE * pct
        billed_wk = r["weekly_rate"] or 0
        diff_wk = prorata_wk - billed_wk
        prorata_mo = prorata_wk * 4.333
        billed_mo = billed_wk * 4.333

        total_prorata_wk += prorata_wk
        total_billed_wk += billed_wk
        total_prorata_mo += prorata_mo
        total_billed_mo += billed_mo

        diff_flag = ""
        if abs(diff_wk) > 1:
            diff_flag = " ^" if diff_wk > 0 else " v"

        print(f"  {r['id']:>3}  {trunc(r['business_name'], 30):<30} "
              f"{trunc(r['address'], 22):<22} "
              f"{sqft:>8,} {pct:>8.4%} {money(prorata_wk):>12} "
              f"{money(billed_wk):>11} {money(diff_wk):>10}{diff_flag} "
              f"{money(prorata_mo):>12} {money(billed_mo):>11}")

    print("  " + "-" * 155)
    diff_total_wk = total_prorata_wk - total_billed_wk
    print(f"  {'':>3}  {'TOTALS':<30} {'':22} "
          f"{'':>8} {'':>9} {money(total_prorata_wk):>12} "
          f"{money(total_billed_wk):>11} {money(diff_total_wk):>10} "
          f"{money(total_prorata_mo):>12} {money(total_billed_mo):>11}")

    print()
    print(f"  ^ = pro-rata is HIGHER than billed (entity owes more than currently billed)")
    print(f"  v = pro-rata is LOWER than billed (entity is being over-billed)")
    print()
    print(f"  Campus total:  {TOTAL_CAMPUS_SQFT:,} SF")
    print(f"  Weekly cost:   {money(CURRENT_WEEKLY_RATE)}")
    print(f"  Net diff/week: {money(diff_total_wk)}")
    print(f"  Net diff/year: {money(diff_total_wk * 52)}")

    # Offer to adjust SF
    print()
    adjust = input("  Adjust square footage for a parcel? (Enter ID, or press Enter to skip): ").strip()
    if adjust:
        try:
            pid = int(adjust)
        except ValueError:
            pause()
            return

        row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
        if not row:
            print(f"  No parcel with ID {pid}.")
            pause()
            return

        print(f"\n  {row['business_name']} — current SF: {(row['sqft'] or 0):,}")
        new_sf = input("  Enter new square footage: ").strip()
        try:
            new_sf = int(new_sf.replace(",", ""))
        except ValueError:
            print("  Invalid number.")
            pause()
            return

        pct = new_sf / TOTAL_CAMPUS_SQFT
        conn.execute("UPDATE parcels SET sqft = ?, pct_campus = ? WHERE id = ?",
                      (new_sf, pct, pid))

        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            """INSERT INTO enforcement_log (parcel_id, timestamp, action, notes)
               VALUES (?,?,?,?)""",
            (pid, now, f"SF updated to {new_sf:,} (was {(row['sqft'] or 0):,})",
             "Pro-rata calculator adjustment"),
        )
        conn.commit()

        new_prorata = CURRENT_WEEKLY_RATE * pct
        print(f"\n  Updated {row['business_name']}:")
        print(f"    SF:              {new_sf:,}")
        print(f"    % of campus:     {pct:.4%}")
        print(f"    Pro-rata weekly:  {money(new_prorata)}")
        print(f"    Pro-rata monthly: {money(new_prorata * 4.333)}")
        print(f"  (Logged at {now})")

    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 11: MARK PACKET SENT
# ═══════════════════════════════════════════════════════════════════════════════

def mark_packet_sent(conn):
    header("MARK PACKET SENT")
    print("  When you mark a packet as sent, the system auto-calculates:")
    print("    - 30-day cure deadline")
    print("    - 45-day lien filing date")
    print("    - 60-day attorney referral date")
    print()

    # Show non-current parcels
    rows = conn.execute(
        "SELECT id, address, business_name, enforcement_step, date_packet_sent "
        "FROM parcels WHERE status != 'CURRENT' ORDER BY id"
    ).fetchall()
    print(f"  {'ID':>3}  {'Address':<24} {'Business':<28} {'Step':<16} {'Sent':<12}")
    print("  " + "-" * 90)
    for r in rows:
        sent = r["date_packet_sent"] or "—"
        print(f"  {r['id']:>3}  {trunc(r['address'], 24):<24} "
              f"{trunc(r['business_name'], 28):<28} "
              f"{trunc(r['enforcement_step'], 16):<16} {sent:<12}")
    print()

    pid = input("  Enter parcel ID to mark as sent: ").strip()
    try:
        pid = int(pid)
    except ValueError:
        print("  Invalid ID.")
        pause()
        return

    row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
    if not row:
        print(f"  No parcel with ID {pid}.")
        pause()
        return

    print(f"\n  Parcel: {row['business_name']} ({row['address']})")

    date_sent = input("  Date sent (YYYY-MM-DD, or Enter for today): ").strip()
    if not date_sent:
        date_sent = datetime.now().strftime("%Y-%m-%d")
    else:
        try:
            datetime.strptime(date_sent, "%Y-%m-%d")
        except ValueError:
            print("  Invalid date format. Use YYYY-MM-DD.")
            pause()
            return

    tracking = input("  Certified mail tracking number (or Enter to skip): ").strip() or None

    deadlines = set_packet_sent(conn, pid, date_sent, tracking)

    print(f"\n  PACKET MARKED AS SENT")
    print(f"  " + "=" * 50)
    print(f"  Parcel:                {row['business_name']}")
    print(f"  Date Sent:             {date_sent}")
    if tracking:
        print(f"  Tracking #:            {tracking}")
    print(f"  30-Day Cure Deadline:  {deadlines['cure_deadline']}")
    print(f"  45-Day Lien Filing:    {deadlines['lien_filing_date']}")
    print(f"  60-Day Attorney Ref:   {deadlines['attorney_referral_date']}")
    print(f"  (All deadlines logged to enforcement timeline)")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 12: UPCOMING DEADLINES
# ═══════════════════════════════════════════════════════════════════════════════

# ANSI color codes for terminal urgency flags
RED = "\033[91m"
YELLOW = "\033[93m"
RESET = "\033[0m"
BOLD = "\033[1m"


def view_deadlines(conn):
    header("UPCOMING DEADLINES")

    today = datetime.now().date()

    # Gather all deadline entries from parcels with packet sent dates
    rows = conn.execute(
        """SELECT id, address, business_name, date_packet_sent,
                  cure_deadline, lien_filing_date, attorney_referral_date,
                  certified_mail_tracking, enforcement_step
           FROM parcels
           WHERE date_packet_sent IS NOT NULL
           ORDER BY cure_deadline"""
    ).fetchall()

    if not rows:
        print("  No packets have been marked as sent yet.")
        print("  Use option 11 (Mark Packet Sent) to record sent dates.")
        pause()
        return

    # Build a flat list of all deadlines for sorting
    deadlines = []
    for r in rows:
        for date_col, label in [
            ("cure_deadline", "30-Day CURE"),
            ("lien_filing_date", "45-Day LIEN"),
            ("attorney_referral_date", "60-Day ATTORNEY"),
        ]:
            if r[date_col]:
                dl_date = datetime.strptime(r[date_col], "%Y-%m-%d").date()
                days_left = (dl_date - today).days
                deadlines.append({
                    "date": r[date_col],
                    "days_left": days_left,
                    "type": label,
                    "business": r["business_name"],
                    "address": r["address"],
                    "id": r["id"],
                    "tracking": r["certified_mail_tracking"],
                    "sent": r["date_packet_sent"],
                })

    deadlines.sort(key=lambda d: d["date"])

    # Print sorted deadline view
    print(f"  {'Date':<12} {'Days':>5} {'Type':<18} {'Business':<28} "
          f"{'Address':<24} {'Tracking #':<24}")
    print("  " + "-" * 115)

    for d in deadlines:
        days = d["days_left"]

        # Color coding
        if days < 0:
            flag = f"{RED}{BOLD} ** OVERDUE **{RESET}"
            prefix = RED
        elif days <= 7:
            flag = f"{RED}{BOLD} ** URGENT **{RESET}"
            prefix = RED
        elif days <= 14:
            flag = f"{YELLOW} * SOON *{RESET}"
            prefix = YELLOW
        else:
            flag = ""
            prefix = ""

        line = (
            f"  {prefix}{d['date']:<12} {days:>5}d {d['type']:<18} "
            f"{trunc(d['business'], 28):<28} "
            f"{trunc(d['address'], 24):<24} "
            f"{trunc(d['tracking'], 24):<24}{RESET}{flag}"
        )
        print(line)

    print("  " + "-" * 115)
    print()

    # Summary counts
    overdue = sum(1 for d in deadlines if d["days_left"] < 0)
    urgent = sum(1 for d in deadlines if 0 <= d["days_left"] <= 7)
    soon = sum(1 for d in deadlines if 7 < d["days_left"] <= 14)
    later = sum(1 for d in deadlines if d["days_left"] > 14)

    if overdue:
        print(f"  {RED}{BOLD}OVERDUE:  {overdue} deadlines past due{RESET}")
    if urgent:
        print(f"  {RED}{BOLD}URGENT:   {urgent} deadlines within 7 days{RESET}")
    if soon:
        print(f"  {YELLOW}SOON:     {soon} deadlines within 14 days{RESET}")
    if later:
        print(f"  OK:       {later} deadlines beyond 14 days")

    print()
    # Show the hard lien deadline
    lien_days = (datetime.strptime(LIEN_DEADLINE, "%Y-%m-%d").date() - today).days
    print(f"  {BOLD}HARD DEADLINE:  {LIEN_DEADLINE}  —  File liens on all non-responders  ({lien_days} days){RESET}")

    # Per-parcel summary
    print()
    print(f"  PACKETS SENT: {len(rows)} parcels")
    print(f"  {'ID':>3}  {'Business':<28} {'Sent':<12} {'Cure':<12} {'Lien':<12} {'Attorney':<12} {'Tracking #':<24}")
    print("  " + "-" * 108)
    for r in rows:
        print(f"  {r['id']:>3}  {trunc(r['business_name'], 28):<28} "
              f"{r['date_packet_sent'] or '—':<12} "
              f"{r['cure_deadline'] or '—':<12} "
              f"{r['lien_filing_date'] or '—':<12} "
              f"{r['attorney_referral_date'] or '—':<12} "
              f"{trunc(r['certified_mail_tracking'], 24):<24}")

    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 13: LENDER / TITLE RESEARCH TRACKER
# ═══════════════════════════════════════════════════════════════════════════════

def lender_research_tracker(conn):
    header("LENDER / TITLE RESEARCH TRACKER")

    print("  Use Shelby County public records to fill in this data:")
    print("    Assessor:  https://www.assessormelvinburgess.com/propertySearch")
    print("    Register:  https://search.register.shelby.tn.us/search/index.php")
    print("    GIS Map:   https://gis.register.shelby.tn.us/")
    print()

    rows = conn.execute(
        """SELECT id, address, business_name, entity_owner, status,
                  county_parcel_id, mailing_address, lender_name, lender_address,
                  deed_of_trust_ref, lender_contact, loan_number, title_company,
                  address_verified, lender_verified
           FROM parcels WHERE status != 'CURRENT' ORDER BY sqft DESC"""
    ).fetchall()

    if not rows:
        print("  No non-current parcels.")
        pause()
        return

    # Summary view
    print(f"  {'ID':>3} {'Business':<28} {'Addr':>4} {'Lndr':>4} "
          f"{'County Parcel':<16} {'Lender':<30} {'DoT Ref':<16}")
    print("  " + "-" * 105)

    verified_addr = 0
    verified_lender = 0
    for r in rows:
        av = "YES" if r["address_verified"] else "---"
        lv = "YES" if r["lender_verified"] else "---"
        if r["address_verified"]:
            verified_addr += 1
        if r["lender_verified"]:
            verified_lender += 1

        cpid = r["county_parcel_id"] or "—"
        lname = r["lender_name"] or "—"
        dot = r["deed_of_trust_ref"] or "—"

        print(f"  {r['id']:>3} {r['business_name']:<28} {av:>4} {lv:>4} "
              f"{cpid:<16} {lname:<30} {dot:<16}")

    print("  " + "-" * 105)
    print(f"  Address verified: {verified_addr}/{len(rows)}  |  "
          f"Lender verified: {verified_lender}/{len(rows)}")
    print()

    # Options
    print("  Options:")
    print("    1. View full detail for a parcel")
    print("    2. Quick-enter lender data for a parcel")
    print("    3. Mark address as verified")
    print("    4. Mark lender as verified")
    print("    0. Back to menu")
    print()

    choice = input("  Choice: ").strip()

    if choice == "1":
        pid = input("  Parcel ID: ").strip()
        try:
            pid = int(pid)
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
        row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
        if not row:
            print(f"  No parcel with ID {pid}.")
            pause()
            return

        print(f"\n  {'=' * 60}")
        print(f"  PARCEL #{row['id']}: {row['business_name']}")
        print(f"  {'=' * 60}")
        print(f"  Address:            {row['address']}")
        print(f"  Mailing Address:    {row['mailing_address'] or '— NOT SET'}")
        print(f"  Entity/Owner:       {row['entity_owner'] or '—'}")
        print(f"  Corporate Target:   {row['corporate_target'] or '—'}")
        print(f"  County Parcel ID:   {row['county_parcel_id'] or '— NOT SET'}")
        print(f"  SqFt:               {(row['sqft'] or 0):,}")
        print(f"  Address Verified:   {'YES' if row['address_verified'] else 'NO'}")
        print()
        print(f"  LENDER / BANKING:")
        print(f"  Lender Name:        {row['lender_name'] or '— NOT SET'}")
        print(f"  Lender Address:     {row['lender_address'] or '— NOT SET'}")
        print(f"  Lender Contact:     {row['lender_contact'] or '— NOT SET'}")
        print(f"  Loan Number:        {row['loan_number'] or '— NOT SET'}")
        print(f"  Deed of Trust Ref:  {row['deed_of_trust_ref'] or '— NOT SET'}")
        print(f"  Title Company:      {row['title_company'] or '— NOT SET'}")
        print(f"  Lender Verified:    {'YES' if row['lender_verified'] else 'NO'}")
        pause()

    elif choice == "2":
        pid = input("  Parcel ID: ").strip()
        try:
            pid = int(pid)
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
        row = conn.execute("SELECT * FROM parcels WHERE id = ?", (pid,)).fetchone()
        if not row:
            print(f"  No parcel with ID {pid}.")
            pause()
            return

        print(f"\n  Quick-enter lender data for: {row['business_name']}")
        print(f"  (Press Enter to skip any field)\n")

        fields = [
            ("county_parcel_id", "County Parcel ID"),
            ("mailing_address", "Mailing Address (for certified mail)"),
            ("lender_name", "Lender / Bank Name"),
            ("lender_address", "Lender Mailing Address"),
            ("lender_contact", "Lender Contact (name/dept)"),
            ("loan_number", "Loan Number"),
            ("deed_of_trust_ref", "Deed of Trust Ref (book/page or instrument #)"),
            ("title_company", "Title Company"),
        ]

        updates = {}
        for field, label in fields:
            current = row[field] or ""
            if current:
                val = input(f"  {label} [{current}]: ").strip()
            else:
                val = input(f"  {label}: ").strip()
            if val:
                updates[field] = val

        if updates:
            set_clause = ", ".join(f"{k} = ?" for k in updates)
            values = list(updates.values()) + [pid]
            conn.execute(f"UPDATE parcels SET {set_clause} WHERE id = ?", values)

            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            fields_updated = ", ".join(updates.keys())
            conn.execute(
                """INSERT INTO enforcement_log (parcel_id, timestamp, action, notes)
                   VALUES (?, ?, ?, ?)""",
                (pid, now, f"Lender research updated: {fields_updated}",
                 "Research tracker entry"),
            )
            conn.commit()
            print(f"\n  Updated {len(updates)} field(s) for {row['business_name']}")
            print(f"  (Logged at {now})")
        else:
            print("  No changes entered.")
        pause()

    elif choice == "3":
        pid = input("  Parcel ID to mark address verified: ").strip()
        try:
            pid = int(pid)
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
        conn.execute("UPDATE parcels SET address_verified = 1 WHERE id = ?", (pid,))
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            """INSERT INTO enforcement_log (parcel_id, timestamp, action, notes)
               VALUES (?, ?, ?, ?)""",
            (pid, now, "Address verified via county records", "Research tracker"),
        )
        conn.commit()
        print("  Address marked as verified.")
        pause()

    elif choice == "4":
        pid = input("  Parcel ID to mark lender verified: ").strip()
        try:
            pid = int(pid)
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
        conn.execute("UPDATE parcels SET lender_verified = 1 WHERE id = ?", (pid,))
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            """INSERT INTO enforcement_log (parcel_id, timestamp, action, notes)
               VALUES (?, ?, ?, ?)""",
            (pid, now, "Lender verified via Register of Deeds", "Research tracker"),
        )
        conn.commit()
        print("  Lender marked as verified.")
        pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MENU 14: GENERATE LENDER NOTIFICATION LETTER
# ═══════════════════════════════════════════════════════════════════════════════

def generate_lender_notification(conn):
    header("GENERATE LENDER NOTIFICATION LETTER")

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  Error: python-docx is not installed.")
        print("  Run: pip install python-docx")
        pause()
        return

    # Show parcels with lender data
    rows = conn.execute(
        """SELECT id, business_name, address, lender_name, lender_address, lender_verified
           FROM parcels
           WHERE status = 'DELINQUENT' AND lender_name IS NOT NULL AND lender_name != ''
           ORDER BY sqft DESC"""
    ).fetchall()

    if not rows:
        print("  No delinquent parcels with lender data found.")
        print("  Use option 13 (Lender Research Tracker) to add lender information first.")
        pause()
        return

    print(f"  Parcels with lender data:")
    print(f"  {'ID':>3} {'Business':<28} {'Lender':<30} {'Verified':>8}")
    print("  " + "-" * 73)
    for r in rows:
        v = "YES" if r["lender_verified"] else "NO"
        print(f"  {r['id']:>3} {r['business_name']:<28} {r['lender_name']:<30} {v:>8}")
    print()

    pid = input("  Parcel ID (or 'all' for batch): ").strip()

    if pid.lower() == "all":
        targets = rows
    else:
        try:
            pid = int(pid)
        except ValueError:
            print("  Invalid ID.")
            pause()
            return
        targets = [r for r in rows if r["id"] == pid]
        if not targets:
            print(f"  No lender data for parcel {pid}.")
            pause()
            return

    generated = []
    for target in targets:
        row = conn.execute("SELECT * FROM parcels WHERE id = ?", (target["id"],)).fetchone()
        sqft = row["sqft"] or 0
        arrears = get_arrears(row)
        today_str = datetime.now().strftime("%B %d, %Y")

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Letterhead
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SECURITY ONE")
        run.bold = True
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("[Letterhead — Address / Phone / Email]")

        doc.add_paragraph("")
        doc.add_paragraph(today_str)
        doc.add_paragraph("")

        # Lender address
        doc.add_paragraph("SENT VIA CERTIFIED MAIL")
        doc.add_paragraph("")
        if row["lender_contact"]:
            doc.add_paragraph(row["lender_contact"])
        doc.add_paragraph(row["lender_name"] or "")
        if row["lender_address"]:
            doc.add_paragraph(row["lender_address"])
        doc.add_paragraph("")

        # Subject
        p = doc.add_paragraph()
        run = p.add_run("RE: NOTICE OF COVENANT DEFAULT AND IMMINENT LIEN — SECURED CREDITOR NOTIFICATION")
        run.bold = True
        run.underline = True
        doc.add_paragraph("")

        # Property identification
        doc.add_paragraph(f"Property Address:    {row['address']}")
        doc.add_paragraph(f"Tenant/Occupant:     {row['business_name']}")
        doc.add_paragraph(f"Record Owner:        {row['entity_owner'] or 'See deed records'}")
        if row["county_parcel_id"]:
            doc.add_paragraph(f"County Parcel ID:    {row['county_parcel_id']}")
        if row["deed_of_trust_ref"]:
            doc.add_paragraph(f"Deed of Trust Ref:   {row['deed_of_trust_ref']}")
        if row["loan_number"]:
            doc.add_paragraph(f"Loan Number:         {row['loan_number']}")
        doc.add_paragraph("")

        # Body
        doc.add_paragraph("Dear Sir or Madam:")
        doc.add_paragraph("")
        doc.add_paragraph(
            "This letter serves as formal notice to you, as the secured lender of record "
            "for the above-referenced property, that the property is in material default of the "
            f"Declaration of Restrictive Covenants recorded on {DECLARATION_DATE} "
            "(the \"Declaration\") governing the Kirby Gate commercial development, Memphis, Tennessee."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "Security One has been designated as the authorized security services provider under "
            "the Declaration. Pursuant to the Declaration, each parcel owner is obligated to fund "
            "its pro-rata share of campus-wide security services."
        )
        doc.add_paragraph("")

        # Arrears
        p = doc.add_paragraph()
        run = p.add_run("DEFAULT AMOUNT:")
        run.bold = True
        doc.add_paragraph(f"  Parcel Square Footage:    {sqft:,} SF of {TOTAL_CAMPUS_SQFT:,} SF campus")
        pct = sqft / TOTAL_CAMPUS_SQFT if sqft else 0
        doc.add_paragraph(f"  Pro-Rata Share:           {pct:.4%}")
        doc.add_paragraph(f"  36-Month Arrears Owed:    {money(arrears)}")
        doc.add_paragraph("")

        # Lien notice
        p = doc.add_paragraph()
        run = p.add_run("NOTICE OF IMMINENT LIEN:")
        run.bold = True
        doc.add_paragraph(
            f"Please be advised that unless the above arrears are cured in full, Security One "
            f"intends to record a Notice of Lien against this property under the Declaration "
            f"on or before {LIEN_DEADLINE}."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "This lien, arising under a recorded Declaration of Restrictive Covenants, will "
            "constitute an encumbrance on the property and may affect the priority, marketability, "
            "and insurability of title. We are providing this notice to allow you, as a secured "
            "creditor, to take whatever action you deem appropriate to protect your interest, "
            "including but not limited to:"
        )
        doc.add_paragraph("  1. Contacting the borrower/owner to demand cure of the covenant default")
        doc.add_paragraph("  2. Exercising any rights under your loan documents relating to covenant compliance")
        doc.add_paragraph("  3. Ensuring that your title insurance covers the covenant lien")
        doc.add_paragraph("")
        doc.add_paragraph(
            "The Declaration provides for lien rights, fee-shifting, and enforcement in "
            "Shelby County, Tennessee. The applicable statute of limitations for contract "
            "enforcement is six (6) years under Tennessee law."
        )
        doc.add_paragraph("")
        doc.add_paragraph(
            "We are available to discuss resolution of this matter and will provide "
            "updates on the status of enforcement proceedings upon request."
        )
        doc.add_paragraph("")
        doc.add_paragraph("Respectfully,")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("____________________________________")
        doc.add_paragraph("Brad")
        doc.add_paragraph("Security One")
        doc.add_paragraph("")
        doc.add_paragraph("cc: Jeff Rosenblum, Esq. (Counsel)")
        doc.add_paragraph(f"    {row['entity_owner'] or 'Property Owner'} (Borrower)")

        # Save
        safe_name = row["business_name"].replace("/", "-").replace(" ", "_")
        safe_lender = (row["lender_name"] or "Lender").replace("/", "-").replace(" ", "_")[:20]
        filename = f"LenderNotice_{safe_name}_{safe_lender}_{datetime.now():%Y%m%d}.docx"
        filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)
        doc.save(filepath)

        # Log
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            """INSERT INTO enforcement_log
               (parcel_id, timestamp, action, sent_via, next_step, attorney, notes)
               VALUES (?,?,?,?,?,?,?)""",
            (target["id"], now,
             f"Lender notification generated: {filename}",
             "USPS Certified (pending)",
             "Send to lender via certified mail",
             "Rosenblum",
             f"Lender: {row['lender_name']}, Arrears: {money(arrears)}"),
        )
        conn.commit()
        generated.append((row["business_name"], row["lender_name"], filepath))

    print(f"\n  Generated {len(generated)} lender notification(s):")
    for biz, lender, fp in generated:
        print(f"    {biz} -> {lender}")
        print(f"      {fp}")
    pause()


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN MENU
# ═══════════════════════════════════════════════════════════════════════════════

def main_menu():
    conn = get_db()

    while True:
        clear()
        days_to_lien = (datetime.strptime(LIEN_DEADLINE, "%Y-%m-%d") - datetime.now()).days
        delinq = conn.execute("SELECT COUNT(*) as c FROM parcels WHERE status='DELINQUENT'").fetchone()["c"]

        print()
        print("  ============================================================")
        print("  KIRBY GATE ENFORCEMENT SYSTEM")
        print("  ============================================================")
        print(f"  {datetime.now():%B %d, %Y %I:%M %p}")
        print(f"  Lien deadline: {LIEN_DEADLINE} ({days_to_lien} days)")
        print(f"  Delinquent parcels: {delinq}")
        print("  ============================================================")
        print()
        print("    1.  View All Parcels")
        print("    2.  View Non-Payers Only")
        print("    3.  Update Parcel Status")
        print("    4.  Log Enforcement Action")
        print("    5.  Generate Demand Letter")
        print("    6.  Settlement Calculator")
        print("    7.  Export to Excel")
        print("    8.  View Enforcement Timeline")
        print("    9.  Dashboard Summary")
        print("   10.  Pro-Rata Calculator")
        print("   11.  Mark Packet Sent")
        print("   12.  Upcoming Deadlines")
        print("  ---- Lender / Title Research ----")
        print("   13.  Lender Research Tracker")
        print("   14.  Generate Lender Notification")
        print("    0.  Exit")
        print()

        choice = input("  Enter choice: ").strip()

        if choice == "1":
            view_all_parcels(conn)
        elif choice == "2":
            view_nonpayers(conn)
        elif choice == "3":
            update_parcel(conn)
        elif choice == "4":
            log_enforcement(conn)
        elif choice == "5":
            generate_demand_letter(conn)
        elif choice == "6":
            settlement_calculator(conn)
        elif choice == "7":
            export_excel(conn)
        elif choice == "8":
            view_timeline(conn)
        elif choice == "9":
            dashboard(conn)
        elif choice == "10":
            prorata_calculator(conn)
        elif choice == "11":
            mark_packet_sent(conn)
        elif choice == "12":
            view_deadlines(conn)
        elif choice == "13":
            lender_research_tracker(conn)
        elif choice == "14":
            generate_lender_notification(conn)
        elif choice == "0":
            print("\n  System closed. All data saved to kirbygate.db.")
            break
        else:
            print("  Invalid choice. Enter 1-14 or 0.")
            pause()

    conn.close()


if __name__ == "__main__":
    main_menu()
